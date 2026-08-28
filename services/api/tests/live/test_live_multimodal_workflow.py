"""Real API/model smoke for the competition-critical multimodal teaching path."""

from __future__ import annotations

import io
import json
import os
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from typing import Any, cast
from uuid import UUID, uuid4

import httpx
import matplotlib.pyplot as plt
import numpy as np
import pytest
from docx import Document
from sqlalchemy import select

from quantum_agent.auth import hash_session_token, issue_opaque_session_token
from quantum_agent.config import Settings
from quantum_agent.credential_vault import build_credential_vault
from quantum_agent.database import create_database_engine, create_session_factory
from quantum_agent.db_models import (
    Course,
    CourseMembership,
    CourseRole,
    CourseStatus,
    CurriculumEdition,
    CurriculumEditionStatus,
    MembershipStatus,
    SessionStatus,
    SystemRole,
    User,
    UserSession,
    UserStatus,
)


@dataclass(frozen=True, slots=True)
class LiveScope:
    course_id: UUID
    edition_id: UUID
    student_token: str
    ta_token: str


JsonObject = dict[str, Any]


def _require_live_model() -> None:
    if os.environ.get("QA_LIVE_MODEL") != "1":
        pytest.skip("run through the Compose live-model service to spend provider calls")


async def _seed_live_scope() -> LiveScope:
    settings = Settings()
    assert settings.database_url.startswith("postgresql+asyncpg://")
    engine = create_database_engine(settings)
    factory = create_session_factory(engine)
    now = datetime.now(UTC)
    try:
        async with factory() as session:
            row = (
                await session.execute(
                    select(Course, CurriculumEdition)
                    .join(CurriculumEdition, CurriculumEdition.course_id == Course.id)
                    .where(CurriculumEdition.status == CurriculumEditionStatus.PUBLISHED)
                    .order_by(
                        CurriculumEdition.published_at.desc(),
                        CurriculumEdition.id.asc(),
                    )
                    .limit(1)
                )
            ).one()
            course, edition = row
            course.status = CourseStatus.ACTIVE
            tokens: dict[CourseRole, str] = {}
            session_ids: list[UUID] = []
            run_id = uuid4()
            for role in (CourseRole.STUDENT, CourseRole.TA):
                user = User(
                    email=f"live-model-{role.value}-{run_id}@quantum-agent.invalid",
                    display_name=f"Live model {role.value}",
                    system_role=SystemRole.USER,
                    status=UserStatus.ACTIVE,
                )
                session.add(user)
                await session.flush()
                session.add(
                    CourseMembership(
                        course_id=course.id,
                        user_id=user.id,
                        role=role,
                        status=MembershipStatus.ACTIVE,
                        joined_at=now,
                    )
                )
                token = issue_opaque_session_token()
                user_session = UserSession(
                    user_id=user.id,
                    session_token_sha256=hash_session_token(token),
                    status=SessionStatus.ACTIVE,
                    expires_at=now + timedelta(hours=2),
                    user_agent="quantum-agent-live-model-test",
                )
                session.add(user_session)
                await session.flush()
                session_ids.append(user_session.id)
                tokens[role] = token
            await session.commit()
        # PRD V3.1 §3.3: the live-model test exercises the real per-session
        # credential path.  Store the deployment USTC_API key in the vault
        # for every seeded session (student AND TA) so authenticated model,
        # vision, and HITL-resume calls route through the session credential
        # (fail-closed would otherwise return 503 for a session with no vault
        # entry).
        vault = build_credential_vault(
            fernet_key=settings.session_vault_key,
            redis_url=settings.effective_redis_url,
        )
        if vault is not None and settings.ustc_api is not None:
            for sid in session_ids:
                await vault.store(sid, settings.ustc_api)
            await vault.close()
        return LiveScope(
            course_id=course.id,
            edition_id=edition.id,
            student_token=tokens[CourseRole.STUDENT],
            ta_token=tokens[CourseRole.TA],
        )
    finally:
        await engine.dispose()


def _derivation_png() -> bytes:
    figure, axis = plt.subplots(figsize=(9, 6), dpi=130)
    figure.patch.set_facecolor("#f4f0e5")
    axis.set_facecolor("#f4f0e5")
    axis.axis("off")
    lines = (
        "Handwritten quantum derivation",
        "Step 1: E = hbar**2*k**2/(2*m)",
        "Step 2: E = hbar**2*k/(2*m)",
        "I cancelled one k here. Is this valid?",
    )
    for index, line in enumerate(lines):
        axis.text(
            0.08 + (0.012 if index % 2 else 0),
            0.84 - index * 0.19,
            line,
            fontsize=18 if index else 20,
            fontfamily="DejaVu Sans Mono",
            fontstyle="italic",
            rotation=(-1.2 if index % 2 else 0.7),
            color="#162e28",
            transform=axis.transAxes,
        )
    output = io.BytesIO()
    figure.savefig(output, format="png", bbox_inches="tight")
    plt.close(figure)
    return output.getvalue()


def _simulation_plot_png() -> bytes:
    time = np.linspace(0, 8, 240)
    excited = np.sin(time / 2) ** 2
    figure, axis = plt.subplots(figsize=(8, 5), dpi=130)
    axis.plot(time, excited, color="#176b55", linewidth=2.4, label="P(excited)")
    axis.set_xlabel("time")
    axis.set_ylabel("population")
    axis.set_title("Two-level Rabi oscillation")
    axis.set_ylim(-0.03, 1.03)
    axis.grid(alpha=0.22)
    axis.legend()
    output = io.BytesIO()
    figure.savefig(output, format="png", bbox_inches="tight")
    plt.close(figure)
    return output.getvalue()


def _lecture_docx() -> bytes:
    document = Document()
    document.add_heading("波函数与归一化", level=1)
    document.add_paragraph("归一化条件要求波函数的总概率为一。")
    document.add_paragraph("先写出学生自己的预测，再检查概率守恒。")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _events(document: str) -> list[tuple[str, JsonObject]]:
    parsed: list[tuple[str, JsonObject]] = []
    for block in document.replace("\r\n", "\n").split("\n\n"):
        if not block.strip():
            continue
        # PRD V3.1 P1-2: the backend now emits ``: keepalive`` comment lines
        # and ``progress`` events between ``workflow.started`` and the
        # terminal event.  Comment-only blocks have no ``event:`` or
        # ``data:`` lines; skip them.  ``progress`` events are informational
        # and carry no terminal contract; ignore them so the terminal event
        # remains the last workflow event in the stream.
        event = "message"
        data_lines: list[str] = []
        for line in block.splitlines():
            if line.startswith(":"):
                continue
            if line.startswith("event:"):
                event = line.removeprefix("event:").strip()
            elif line.startswith("data:"):
                data_lines.append(line.removeprefix("data:").lstrip())
        if event == "progress" or not data_lines:
            continue
        payload = json.loads("\n".join(data_lines))
        assert isinstance(payload, dict)
        parsed.append((event, cast(JsonObject, payload)))
    return parsed


async def _confirm_if_needed(
    client: httpx.AsyncClient,
    *,
    base: str,
    headers: dict[str, str],
    attachment: JsonObject,
) -> JsonObject:
    extraction = attachment.get("extraction")
    assert isinstance(extraction, dict)
    status = extraction.get("status")
    assert status in {"succeeded", "needs_confirmation", "confirmed"}
    if status != "needs_confirmation":
        return attachment
    ambiguities = extraction.get("ambiguities")
    assert isinstance(ambiguities, list)
    resolutions: dict[str, str] = {}
    for raw in ambiguities:
        assert isinstance(raw, dict)
        ambiguity_id = raw.get("ambiguity_id")
        if not isinstance(ambiguity_id, str):
            continue
        candidates = raw.get("candidates")
        selected = "confirmed exactly as transcribed"
        if isinstance(candidates, list) and candidates and isinstance(candidates[0], dict):
            candidate = candidates[0].get("value")
            if isinstance(candidate, str) and candidate.strip():
                selected = candidate
        resolutions[ambiguity_id] = selected
    response = await client.post(
        f"{base}/{attachment['id']}/confirm",
        headers=headers,
        json={
            "extraction_id": extraction["id"],
            "decision": "accept",
            "ambiguity_resolutions": resolutions,
        },
    )
    assert response.status_code == 200, response.text[:1000]
    confirmed = cast(JsonObject, response.json())
    assert confirmed["extraction"]["status"] == "confirmed"
    return confirmed


async def _teaching_turn(
    client: httpx.AsyncClient,
    *,
    scope: LiveScope,
    headers: dict[str, str],
    body: JsonObject,
) -> tuple[JsonObject, bool]:
    path = (
        f"/api/v1/courses/{scope.course_id}/editions/{scope.edition_id}"
        "/teaching/turns/stream"
    )
    response = await client.post(path, headers=headers, json=body)
    assert response.status_code == 200, response.text[:1000]
    events = _events(response.text)
    assert events, f"no workflow events parsed from stream: {response.text[:500]!r}"
    first_event = events[0][0]
    terminal_event, outcome = events[-1]
    assert first_event == "workflow.started", events
    assert terminal_event in {"workflow.completed", "workflow.interrupted"}, events
    return outcome, terminal_event == "workflow.interrupted"


async def _approve_interrupt(
    client: httpx.AsyncClient,
    *,
    scope: LiveScope,
    conversation_id: str,
) -> JsonObject:
    path = (
        f"/api/v1/courses/{scope.course_id}/editions/{scope.edition_id}"
        f"/teaching/threads/{conversation_id}"
    )
    headers = {"Authorization": f"Bearer {scope.ta_token}"}
    inspected = await client.get(f"{path}/interrupt", headers=headers)
    assert inspected.status_code == 200, inspected.text[:1000]
    assert inspected.json()["conversation_id"] == conversation_id
    resumed = await client.post(f"{path}/resume", headers=headers, json={"action": "approve"})
    assert resumed.status_code == 200, resumed.text[:1000]
    result = cast(JsonObject, resumed.json())
    assert result["conversation_id"] == conversation_id
    return result


async def _confirm_transcription_interrupt(
    client: httpx.AsyncClient,
    *,
    scope: LiveScope,
    conversation_id: str,
    confirmed_student_attempt: str,
) -> JsonObject:
    path = (
        f"/api/v1/courses/{scope.course_id}/editions/{scope.edition_id}"
        f"/teaching/threads/{conversation_id}"
    )
    headers = {"Authorization": f"Bearer {scope.student_token}"}
    inspected = await client.get(f"{path}/interrupt", headers=headers)
    assert inspected.status_code == 200, inspected.text[:1000]
    pause = cast(JsonObject, inspected.json())
    assert pause["conversation_id"] == conversation_id
    assert pause["interrupt"]["reasons"] == ["ambiguous_transcription"]
    assert pause["interrupt"]["student_allowed_actions"] == ["confirm_transcription"]
    resumed = await client.post(
        f"{path}/resume",
        headers=headers,
        json={
            "action": "confirm_transcription",
            "confirmed_student_attempt": confirmed_student_attempt,
        },
    )
    assert resumed.status_code == 200, resumed.text[:1000]
    result = cast(JsonObject, resumed.json())
    assert result["conversation_id"] == conversation_id
    return result


@pytest.mark.live_model
@pytest.mark.asyncio
async def test_real_multimodal_derivation_document_plot_hitl_and_trace() -> None:
    _require_live_model()
    scope = await _seed_live_scope()
    api_base = os.environ.get("QUANTUM_API_BASE_URL", "http://api:8000").rstrip("/")
    student_headers = {"Authorization": f"Bearer {scope.student_token}"}
    attachment_base = (
        f"/api/v1/courses/{scope.course_id}/editions/{scope.edition_id}/attachments"
    )
    timeout = httpx.Timeout(330.0, connect=20.0)
    async with httpx.AsyncClient(base_url=api_base, timeout=timeout) as client:
        derivation_upload = await client.post(
            attachment_base,
            headers=student_headers,
            files={"file": ("handwritten-derivation.png", _derivation_png(), "image/png")},
        )
        assert derivation_upload.status_code in {200, 201}, derivation_upload.text[:1000]
        derivation = cast(JsonObject, derivation_upload.json())
        derivation_evidence = derivation["extraction"]["evidence"]
        assert derivation_evidence["derivation_steps"]
        assert derivation_evidence["original_file_reference"] == f"attachment:{derivation['id']}"
        derivation_needs_confirmation = (
            derivation["extraction"]["status"] == "needs_confirmation"
        )

        derivation_result, interrupted = await _teaching_turn(
            client,
            scope=scope,
            headers=student_headers,
            body={
                "mode": "review_derivations",
                "message": "检查这个动能本征值推导，定位第一个有后果的错误。",
                "attachment_ids": [derivation["id"]],
            },
        )
        if derivation_needs_confirmation:
            assert interrupted is True
            derivation_result = await _confirm_transcription_interrupt(
                client,
                scope=scope,
                conversation_id=str(derivation_result["conversation_id"]),
                confirmed_student_attempt=(
                    "E = hbar**2*k**2/(2*m)\n"
                    "E = hbar**2*k/(2*m)"
                ),
            )
        elif interrupted:
            derivation_result = await _approve_interrupt(
                client,
                scope=scope,
                conversation_id=str(derivation_result["conversation_id"]),
            )
        assert derivation_result["diagnosis"]["first_error"] is not None
        assert derivation_result["scientific_results"]
        assert derivation_result["scientific_results"][0]["kind"] == "symbolic_equivalence"
        assert derivation_result["evidence_packet"]["evidence"]
        assert derivation_result["validation"]["passed"] is True

        first_citation = derivation_result["evidence_packet"]["evidence"][0]
        source = await client.get(
            (
                f"/api/v1/courses/{scope.course_id}/editions/{scope.edition_id}/sources/"
                f"{first_citation['document_version_id']}/original"
            ),
            headers={**student_headers, "Range": "bytes=0-1023"},
        )
        assert source.status_code in {200, 206}
        assert source.content
        assert source.headers.get("etag", "").startswith('"sha256-')

        document_upload = await client.post(
            attachment_base,
            headers=student_headers,
            files={
                "file": (
                    "lecture-note.docx",
                    _lecture_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert document_upload.status_code in {200, 201}, document_upload.text[:1000]
        document_attachment = await _confirm_if_needed(
            client,
            base=attachment_base,
            headers=student_headers,
            attachment=cast(JsonObject, document_upload.json()),
        )
        document_evidence = document_attachment["extraction"]["evidence"]
        assert document_evidence["units"]
        assert document_evidence["fallback_chain"][0]["method"] == "native"

        document_result, interrupted = await _teaching_turn(
            client,
            scope=scope,
            headers=student_headers,
            body={
                "mode": "learn_concepts",
                "message": "结合课程证据，说明这份讲义中的波函数归一化与概率解释如何连接。",
                "attachment_ids": [document_attachment["id"]],
            },
        )
        if interrupted:
            document_result = await _approve_interrupt(
                client,
                scope=scope,
                conversation_id=str(document_result["conversation_id"]),
            )
        assert document_result["interpretation"]["relevant_concepts"]
        assert document_result["evidence_packet"]["evidence"]

        plot_upload = await client.post(
            attachment_base,
            headers=student_headers,
            files={"file": ("rabi-plot.png", _simulation_plot_png(), "image/png")},
        )
        assert plot_upload.status_code in {200, 201}, plot_upload.text[:1000]
        plot_attachment = await _confirm_if_needed(
            client,
            base=attachment_base,
            headers=student_headers,
            attachment=cast(JsonObject, plot_upload.json()),
        )
        plot_evidence = plot_attachment["extraction"]["evidence"]
        assert plot_evidence["plot_axes"]

        plot_result, interrupted = await _teaching_turn(
            client,
            scope=scope,
            headers=student_headers,
            body={
                "mode": "run_experiments",
                "message": "先检查数值不变量，再解释上传的 Rabi 振荡图。",
                "attachment_ids": [plot_attachment["id"]],
                "scientific_request": {
                    "kind": "two_level_simulation",
                    "initial_state": [{"real": 1, "imag": 0}, {"real": 0, "imag": 0}],
                    "rabi_frequency": 1.0,
                    "detuning": 0.0,
                    "duration": 8.0,
                    "steps": 400,
                    "absolute_tolerance": 1e-7,
                },
            },
        )
        if interrupted:
            plot_result = await _approve_interrupt(
                client,
                scope=scope,
                conversation_id=str(plot_result["conversation_id"]),
            )
        simulation = next(
            item
            for item in plot_result["scientific_results"]
            if item["kind"] == "two_level_simulation"
        )
        assert simulation["status"] == "pass"
        assert simulation["visualization"] is not None

        interrupt, was_interrupted = await _teaching_turn(
            client,
            scope=scope,
            headers=student_headers,
            body={
                "mode": "learn_concepts",
                "message": "@TA 请检查我对波函数概率解释的理解。",
                "student_attempt": "我认为波函数本身就是直接可观测概率。",
            },
        )
        assert was_interrupted is True
        assert "ta_requested" in interrupt["interrupt"]["reasons"]
        conversation_id = str(interrupt["conversation_id"])
        resumed = await _approve_interrupt(
            client,
            scope=scope,
            conversation_id=conversation_id,
        )
        assert resumed["conversation_id"] == conversation_id

        trace_base = (
            f"/api/v1/courses/{scope.course_id}/editions/{scope.edition_id}"
            "/teacher/agent-traces"
        )
        ta_headers = {"Authorization": f"Bearer {scope.ta_token}"}
        trace_page = await client.get(trace_base, headers=ta_headers)
        assert trace_page.status_code == 200, trace_page.text[:1000]
        matching = next(
            item
            for item in trace_page.json()["items"]
            if item["conversation_id"] == conversation_id
        )
        detail = await client.get(f"{trace_base}/{matching['id']}", headers=ta_headers)
        assert detail.status_code == 200, detail.text[:1000]
        trace = detail.json()
        assert trace["evidence_bundle"] is not None
        assert trace["diagnosis"] is not None
        assert trace["release_decision"] is not None
        assert trace["hitl_events"]
